from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import os
import tempfile
from PIL import Image
import io


def validate_and_convert_image(image_path):
    """
    Valide et convertit une image pour qu'elle soit compatible avec python-docx

    Args:
        image_path (str): Chemin vers l'image à valider

    Returns:
        str: Chemin vers l'image convertie (ou originale si compatible)
        None: Si l'image ne peut pas être traitée
    """
    if not os.path.exists(image_path):
        print(f"[WARNING] Image non trouvée: {image_path}")
        return None

    try:
        with Image.open(image_path) as img:
            # Vérifier le format et le mode de l'image
            print(f"[IMAGE] Image détectée: {img.format} {img.mode} {img.size}")

            # Formats supportés par python-docx : JPEG, PNG, BMP, TIFF
            supported_formats = ['JPEG', 'PNG', 'BMP', 'TIFF']

            # Modes problématiques qui nécessitent une conversion
            problematic_modes = ['CMYK', 'P', 'L', 'LA']

            # Si l'image est déjà dans un format supporté et un mode compatible
            if (img.format in supported_formats and
                img.mode not in problematic_modes and
                img.mode in ['RGB', 'RGBA']):
                print(f"[OK] Image compatible: {img.format} {img.mode}")
                return image_path

            # Conversion nécessaire
            print(f"[CONVERT] Conversion nécessaire: {img.format} {img.mode} → PNG RGB")

            # Convertir l'image en mode RGB
            if img.mode in ['CMYK', 'L', 'LA']:
                # Conversion spéciale pour CMYK
                if img.mode == 'CMYK':
                    img = img.convert('RGB')
                    print("   CMYK → RGB")
                # Conversion pour niveaux de gris
                elif img.mode in ['L', 'LA']:
                    img = img.convert('RGB')
                    print("   Niveaux de gris → RGB")
            elif img.mode == 'P':
                # Conversion pour palette
                img = img.convert('RGBA')
                print("   Palette → RGBA")
            elif img.mode not in ['RGB', 'RGBA']:
                # Conversion générique vers RGB
                img = img.convert('RGB')
                print(f"   {img.mode} → RGB")

            # Créer un fichier temporaire pour l'image convertie
            temp_dir = tempfile.gettempdir()
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            temp_path = os.path.join(temp_dir, f"{base_name}_converted.png")

            # Sauvegarder l'image convertie
            img.save(temp_path, 'PNG', optimize=True)
            print(f"[OK] Image convertie sauvegardée: {temp_path}")

            return temp_path

    except Exception as e:
        print(f"[ERROR] Erreur lors de la validation/conversion de l'image: {e}")
        return None


def add_image_to_document(doc, image_path, width_inches=5.5):
    """
    Ajoute une image au document Word de manière sécurisée

    Args:
        doc: Document Word
        image_path (str): Chemin vers l'image
        width_inches (float): Largeur de l'image en pouces

    Returns:
        bool: True si l'image a été ajoutée avec succès, False sinon
    """
    if not image_path or not os.path.exists(image_path):
        print("[WARNING] Aucune image à ajouter ou image inexistante")
        return False

    try:
        # Valider et convertir l'image si nécessaire
        processed_image_path = validate_and_convert_image(image_path)

        if processed_image_path is None:
            print("[ERROR] Impossible de traiter l'image pour l'export Word")
            return False

        # Ajouter l'image au document
        doc.add_picture(processed_image_path, width=Inches(width_inches))
        print(f"[OK] Image ajoutée au document Word: {os.path.basename(image_path)}")

        # Nettoyer le fichier temporaire si c'était une conversion
        if processed_image_path != image_path and os.path.exists(processed_image_path):
            try:
                os.remove(processed_image_path)
                print("[CLEAN] Fichier temporaire nettoyé")
            except Exception as e:
                print(f"[WARNING] Impossible de supprimer le fichier temporaire: {e}")

        return True

    except Exception as e:
        print(f"[ERROR] Erreur lors de l'ajout de l'image au document: {e}")
        return False


def export_to_word(text_lines, confidences, output_path, image_path=None, method=None):
    """
    Exporte les résultats OCR vers un document Word avec gestion robuste des images

    Args:
        text_lines (list): Lignes de texte extraites par OCR
        confidences (list): Scores de confiance correspondants
        output_path (str): Chemin de sortie du fichier Word
        image_path (str, optional): Chemin vers l'image source
        method (str, optional): Méthode OCR utilisée
    """
    print(f"[EXPORT] Export Word vers: {output_path}")

    try:
        doc = Document()
        doc.add_heading('Résultat OCR', 0)

        # Ajouter des informations sur la méthode utilisée
        if method:
            info_paragraph = doc.add_paragraph()
            info_run = info_paragraph.add_run(f"Méthode OCR utilisée: {method.upper()}")
            info_run.bold = True
            doc.add_paragraph()  # Ligne vide

        # Ajouter l'image de manière sécurisée
        image_added = False
        if image_path is not None:
            print(f"[IMAGE] Tentative d'ajout de l'image: {image_path}")
            image_added = add_image_to_document(doc, image_path, width_inches=5.5)

            if image_added:
                doc.add_paragraph()  # Ligne vide après l'image

        # Ajouter le texte extrait
        if image_added:
            doc.add_paragraph("Texte extrait de l'image :")
        else:
            doc.add_paragraph("Texte extrait :")

        doc.add_paragraph()  # Ligne vide

        # Ajouter les lignes de texte avec coloration selon la confiance
        if text_lines:
            for line, conf in zip(text_lines, confidences):
                if line.strip():  # Ignorer les lignes vides
                    p = doc.add_paragraph()
                    run = p.add_run(line)

                    # Coloration selon le niveau de confiance
                    if conf < 50:
                        run.font.color.rgb = RGBColor(255, 0, 0)    # Rouge pour très faible confiance
                    elif conf < 70:
                        run.font.color.rgb = RGBColor(255, 165, 0)  # Orange pour faible confiance
                    elif conf < 85:
                        run.font.color.rgb = RGBColor(255, 255, 0)  # Jaune pour confiance moyenne
                    # Pas de coloration pour confiance élevée (>= 85)
        else:
            # Aucun texte détecté
            p = doc.add_paragraph()
            run = p.add_run("Aucun texte détecté dans l'image.")
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gris

        # Ajouter des statistiques de confiance
        if confidences and len(confidences) > 0:
            doc.add_paragraph()  # Ligne vide
            stats_paragraph = doc.add_paragraph()
            avg_conf = sum(confidences) / len(confidences)
            min_conf = min(confidences)
            max_conf = max(confidences)

            stats_text = f"Statistiques de confiance - Moyenne: {avg_conf:.1f}%, Min: {min_conf:.1f}%, Max: {max_conf:.1f}%"
            stats_run = stats_paragraph.add_run(stats_text)
            stats_run.italic = True
            stats_run.font.color.rgb = RGBColor(100, 100, 100)

        # Sauvegarder le document
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        print(f"[OK] Document Word exporté avec succès: {output_path}")

        return True

    except Exception as e:
        print(f"[ERROR] Erreur lors de l'export Word: {e}")
        import traceback
        traceback.print_exc()
        return False