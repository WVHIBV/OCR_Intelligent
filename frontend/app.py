"""
Interface Streamlit pour l'application OCR Intelligent
"""
import os
import sys
import glob
import shutil
from pathlib import Path

# Configuration de l'environnement
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"
os.environ["CUDA_VISIBLE_DEVICES"] = ""
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"

import warnings
warnings.filterwarnings("ignore")

import streamlit as st
import fitz
from PIL import Image
from streamlit.components.v1 import html
import docx
import zipfile

# Configuration du chemin backend
current_dir = Path(__file__).parent
backend_dir = current_dir.parent
sys.path.insert(0, str(backend_dir))

from backend.main import run_all_ocr_methods
from backend.export import export_to_word
from backend.preprocessing import detect_text_zones

def create_simple_word_document(zone_ocr_results, original_image_path):
    """
    Crée un document Word simple avec juste le texte réorganisé
    """
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    try:
        # Créer le document
        doc = docx.Document()

        # Titre principal
        title = doc.add_heading('📄 TEXTE EXTRAIT ET RÉORGANISÉ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informations générales
        doc.add_paragraph(f"Document source: {os.path.basename(original_image_path)}")
        doc.add_paragraph(f"Date de traitement: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

        # Compter les zones réussies
        successful_zones = [z for z in zone_ocr_results.values() 
                           if isinstance(z, dict) and "error" not in z]
        doc.add_paragraph(f"Zones traitées: {len(successful_zones)}")

        # Ligne de séparation
        doc.add_paragraph("─" * 50)

        # Récupérer l'ordre de lecture intelligent
        reading_order = zone_ocr_results.get("reading_order", [])
        
        if reading_order:
            # Filtrer les zones valides
            valid_zones = {zone_id: zone_data for zone_id, zone_data in zone_ocr_results.items() 
                          if isinstance(zone_data, dict) and "error" not in zone_data}
            
            # Réorganiser le texte
            reorganized_text = reorganize_text_by_reading_order(valid_zones, reading_order)
            
            if reorganized_text and reorganized_text != "Aucun texte valide détecté dans les zones":
                # Titre de section
                doc.add_heading('📖 TEXTE RÉORGANISÉ SELON L\'ORDRE DE LECTURE INTELLIGENT', level=1)
                
                # Ajouter le texte réorganisé
                text_para = doc.add_paragraph()
                text_run = text_para.add_run(reorganized_text)
                text_run.font.name = 'Arial'
                text_run.font.size = docx.shared.Pt(11)
                
                # Statistiques
                doc.add_paragraph("─" * 30)
                doc.add_heading('📊 Statistiques', level=2)
                
                if successful_zones:
                    avg_conf = sum(z["confidence"] for z in successful_zones) / len(successful_zones)
                    doc.add_paragraph(f"Confiance moyenne: {avg_conf:.1f}%")
                    
                    total_chars = sum(len(z.get("best_text", "")) for z in successful_zones)
                    doc.add_paragraph(f"Caractères extraits: {total_chars:,}")
            else:
                doc.add_paragraph("⚠️ Aucun texte valide n'a pu être extrait des zones détectées")
        else:
            doc.add_paragraph("⚠️ Aucun ordre de lecture intelligent disponible")

        # Sauvegarder le document
        output_path = os.path.join("output", f"{os.path.splitext(os.path.basename(original_image_path))[0]}_texte_reorganise.docx")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        print(f"[OK] Document Word simple créé: {output_path}")
        return output_path

    except Exception as e:
        print(f"[ERROR] Erreur lors de la création du document Word simple: {e}")
        return None

def reorganize_text_by_reading_order(zone_results, reading_order):
    """
    Réorganise le texte extrait selon l'ordre de lecture intelligent
    
    Args:
        zone_results (dict): Résultats des zones avec leurs données
        reading_order (list): Liste ordonnée des IDs de zones
    
    Returns:
        str: Texte réorganisé selon l'ordre de lecture
    """
    if not zone_results or not reading_order:
        return "Aucun texte à réorganiser"
    
    reorganized_text = []
    
    for zone_id in reading_order:
        if zone_id in zone_results:
            zone_data = zone_results[zone_id]
            best_text = zone_data.get("best_text", "").strip()
            
            # Filtrer les textes vides ou qui contiennent des messages d'erreur
            if (best_text and 
                best_text != "Aucun texte détecté avec toutes les configurations" and
                not best_text.startswith("Aucun texte détecté") and
                len(best_text) > 1):  # Au moins 2 caractères
                reorganized_text.append(best_text)
    
    if not reorganized_text:
        return "Aucun texte valide détecté dans les zones"
    
    # Ajouter des séparateurs pour améliorer la lisibilité
    if len(reorganized_text) > 1:
        # Ajouter des espaces entre les zones pour une meilleure séparation
        return "\n\n".join(reorganized_text)
    else:
        return "\n".join(reorganized_text)

def _create_zones_word_document(zone_ocr_results, original_image_path):
    """
    Crée un document Word avec les résultats OCR de toutes les zones
    """
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Créer le document
    doc = docx.Document()

    # Titre principal
    title = doc.add_heading('Résultats OCR par Zones', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Informations générales
    doc.add_paragraph(f"Document source: {os.path.basename(original_image_path)}")

    from datetime import datetime
    doc.add_paragraph(f"Date de traitement: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # Compter les zones réussies (ignorer les clés spéciales)
    successful_zones = [z for z in zone_ocr_results.values() 
                       if isinstance(z, dict) and "error" not in z]
    doc.add_paragraph(f"Zones traitées avec succès: {len(successful_zones)}/{len(zone_ocr_results)}")

    # Ajouter une ligne de séparation
    doc.add_paragraph("─" * 50)

    # Ajouter le texte réorganisé selon l'ordre de lecture intelligent
    reading_order = zone_ocr_results.get("reading_order", [])
    if reading_order:
        doc.add_heading('📄 TEXTE RÉORGANISÉ SELON L\'ORDRE DE LECTURE INTELLIGENT', level=1)
        
        # Filtrer les zones valides
        valid_zones = {zone_id: zone_data for zone_id, zone_data in zone_ocr_results.items() 
                      if isinstance(zone_data, dict) and "error" not in zone_data}
        
        # Réorganiser le texte
        reorganized_text = reorganize_text_by_reading_order(valid_zones, reading_order)
        
        if reorganized_text and reorganized_text != "Aucun texte à réorganiser":
            text_para = doc.add_paragraph()
            text_run = text_para.add_run(reorganized_text)
            text_run.font.name = 'Arial'
            text_run.font.size = docx.shared.Pt(11)
        else:
            doc.add_paragraph("(Aucun texte à réorganiser)")
        
        doc.add_paragraph("─" * 50)

    # Traiter chaque zone - gérer le tri des clés mixtes (str et int)
    def sort_key(zone_id):
        if isinstance(zone_id, str):
            # Essayer de convertir en int, sinon garder la string
            try:
                return (0, int(zone_id))  # 0 pour les entiers
            except ValueError:
                return (1, zone_id)  # 1 pour les strings
        else:
            return (0, zone_id)  # 0 pour les entiers
    
    for zone_id in sorted(zone_ocr_results.keys(), key=sort_key):
        zone_data = zone_ocr_results[zone_id]
        
        # Ignorer les clés spéciales qui ne sont pas des zones (comme "reading_order")
        if not isinstance(zone_data, dict) or zone_id == "reading_order":
            continue

        if "error" in zone_data:
            # Zone en erreur
            heading = doc.add_heading(f'Zone {zone_id} - ERREUR', level=1)
            doc.add_paragraph(f"Erreur: {zone_data['error']}")
        else:
            # Zone réussie
            best_method = zone_data["best_method"]
            confidence = zone_data["confidence"]

            # Titre de la zone
            heading = doc.add_heading(f'Zone {zone_id} - {best_method.upper()}', level=1)

            # Informations sur la zone
            info_para = doc.add_paragraph()
            info_para.add_run("Confiance: ").bold = True
            info_para.add_run(f"{confidence:.1f}%")

            zone_info = zone_data["zone_info"]
            coords = zone_info["coordinates"]
            info_para.add_run(" | Position: ").bold = True
            info_para.add_run(f"({coords['x']}, {coords['y']}) - {coords['width']}×{coords['height']}px")

            # Ajouter l'image de la zone si possible
            try:
                if os.path.exists(zone_info["path"]):
                    doc.add_paragraph().add_run().add_picture(zone_info["path"], width=Inches(3))
            except Exception:
                pass  # Ignorer si l'image ne peut pas être ajoutée

            # Texte extrait
            text_heading = doc.add_heading('Texte extrait:', level=2)

            best_text = zone_data["best_text"]
            if best_text.strip():
                # Ajouter le texte dans un paragraphe avec style
                text_para = doc.add_paragraph()
                text_run = text_para.add_run(best_text)
                text_run.font.name = 'Courier New'

                # Encadrer le texte
                text_para.style = 'Quote'
            else:
                doc.add_paragraph("(Aucun texte détecté)")

        # Séparateur entre les zones
        doc.add_paragraph("─" * 30)

    # Résumé final
    if successful_zones:
        doc.add_heading('Résumé', level=1)

        # Statistiques globales
        total_confidence = sum(z["confidence"] for z in successful_zones)
        avg_confidence = total_confidence / len(successful_zones)
        total_chars = sum(len(z["best_text"]) for z in successful_zones)

        summary_para = doc.add_paragraph()
        summary_para.add_run("Confiance moyenne: ").bold = True
        summary_para.add_run(f"{avg_confidence:.1f}%\n")
        summary_para.add_run("Total de caractères extraits: ").bold = True
        summary_para.add_run(f"{total_chars:,}")

        # Texte complet consolidé
        doc.add_heading('Texte complet (toutes zones)', level=2)
        full_text = "\n\n".join(z["best_text"] for z in successful_zones if z["best_text"].strip())

        if full_text:
            full_para = doc.add_paragraph(full_text)
            full_para.style = 'Quote'
        else:
            doc.add_paragraph("(Aucun texte consolidé disponible)")

    # Sauvegarder le document
    os.makedirs("output", exist_ok=True)
    base_name = os.path.splitext(os.path.basename(original_image_path))[0]
    word_filename = f"result_zones_{base_name}.docx"
    word_path = os.path.join("output", word_filename)

    doc.save(word_path)
    return word_path


def clear_output_directory():
    """
    Nettoie le dossier output à chaque relancement de l'application
    Supprime tous les fichiers .docx générés précédemment
    """
    output_dir = Path("output")

    if output_dir.exists():
        try:
            # Compter les fichiers avant nettoyage
            files_before = list(output_dir.glob("*.docx"))
            files_count = len(files_before)

            if files_count > 0:
                # Supprimer tous les fichiers .docx
                cleaned_count = 0
                for file_path in files_before:
                    try:
                        file_path.unlink()
                        cleaned_count += 1
                    except Exception as e:
                        st.warning(f"Impossible de supprimer {file_path}: {e}")

                if cleaned_count > 0:
                    st.info(f"Dossier output nettoyé: {cleaned_count} fichier(s) supprimé(s)")
            else:
                st.info("Dossier output déjà propre")

        except Exception as e:
            st.error(f"Erreur lors du nettoyage du dossier output: {e}")
    else:
        # Créer le dossier s'il n'existe pas
        output_dir.mkdir(exist_ok=True)
        st.info("Dossier output créé")


# [CLEAN] Nettoyage automatique du dossier output au démarrage
if 'output_cleaned' not in st.session_state:
    clear_output_directory()
    st.session_state.output_cleaned = True

st.set_page_config(page_title="OCR Intelligent", layout="wide")
st.image("frontend/safran_logo.png", width=250)
with open("frontend/custom_style.html") as f:
    html(f.read(), height=0)
st.markdown("<h1> OCR Intelligent </h1>", unsafe_allow_html=True)

uploaded_file = st.file_uploader(" Téléversez une image ou un PDF", type=["png", "jpg", "jpeg", "pdf"])

if uploaded_file:
    os.makedirs("images", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    file_path = os.path.join("images", uploaded_file.name)

    if uploaded_file.type == "application/pdf":
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        st.warning(f"PDF détecté, traitement de {len(doc)} pages.")
        image_paths = []
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            image_path = os.path.join("images", f"page_{i+1}.png")
            pix.save(image_path)
            image_paths.append(image_path)
    else:
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        image_paths = [file_path]

    # --- Configuration principale du traitement
    st.markdown("---")
    st.markdown("### ⚙️ Configuration du traitement")

    # Sélection du type de document (principal)
    document_types = {
        "default": "🔧 Standard (OCR classique)",
        "facture": "📄 Facture",
        "formulaire": "📝 Formulaire",
        "journal": "📰 Journal/Magazine",
        "manuscrit": "✍️ Manuscrit",
        "tableau": "📊 Tableau",
        "photo": "📸 Photo de document"
    }

    selected_doc_type = st.selectbox(
        "Type de document",
        options=list(document_types.keys()),
        index=1,  # Par défaut sur "facture"
        format_func=lambda x: document_types[x],
        help="Sélectionnez le type de document pour optimiser le traitement OCR"
    )

    # Options avancées dans un expander pour ne pas surcharger
    with st.expander("🔧 Options avancées", expanded=False):
        col_zone_opt, col_ocr_opt = st.columns(2)

        with col_zone_opt:
            enable_zone_detection = st.checkbox(
                "Isolation des zones de texte",
                value=(selected_doc_type != "default"),
                help="Détecte et isole automatiquement les zones de texte pour une meilleure précision",
                disabled=(selected_doc_type == "default")
            )

            # Option pour le système intelligent
            if enable_zone_detection:
                use_intelligent_detection = st.checkbox(
                    "🧠 Détection intelligente (NOUVEAU)",
                    value=True,
                    help="Utilise le nouveau système de détection avec classification sémantique et ordre de lecture intelligent"
                )

        with col_ocr_opt:
            if selected_doc_type == "default":
                st.info("Mode standard : OCR classique sur l'image complète")
            else:
                ocr_mode = st.radio(
                    "Mode OCR",
                    ["Image complète", "Zones isolées", "Les deux"],
                    index=2,
                    help="Choisissez comment appliquer l'OCR"
                )

    # --- Traitement principal
    st.markdown("---")

    # Détection des zones si activée et pas en mode standard
    zone_results = None
    if selected_doc_type != "default" and enable_zone_detection:
        # Déterminer le type de détection à utiliser
        detection_type = "intelligente" if use_intelligent_detection else "classique"

        with st.spinner(f"🔍 Détection {detection_type} des zones de texte en cours..."):
            zone_results = detect_text_zones(
                image_paths[0],
                document_type=selected_doc_type,
                use_intelligent_detection=use_intelligent_detection
            )

        if zone_results and zone_results["success"]:
            total_zones = zone_results['total_zones']

            # Affichage différencié selon le type de détection
            if use_intelligent_detection:
                st.success(f"🧠 {total_zones} zones intelligemment détectées et classifiées")

                # Afficher les types de zones détectées
                if "zone_types" in zone_results:
                    zone_types = zone_results["zone_types"]
                    if zone_types:
                        st.info("📊 **Types de zones détectées :**")
                        cols = st.columns(min(4, len(zone_types)))
                        for i, (zone_type, count) in enumerate(zone_types.items()):
                            with cols[i % len(cols)]:
                                # Émojis pour les types
                                type_emojis = {
                                    "header": "🏷️", "price": "💰", "date": "📅",
                                    "address": "🏠", "reference": "📄", "paragraph": "📝",
                                    "signature": "✍️", "footer": "📋", "unknown": "❓"
                                }
                                emoji = type_emojis.get(zone_type, "📄")
                                st.metric(f"{emoji} {zone_type.title()}", count)

                # Afficher l'ordre de lecture
                if "reading_order" in zone_results and zone_results["reading_order"]:
                    with st.expander("📖 Ordre de lecture intelligent"):
                        st.write("Les zones seront traitées dans cet ordre logique :")
                        order_text = " → ".join([f"Zone {i}" for i in zone_results["reading_order"][:10]])
                        if len(zone_results["reading_order"]) > 10:
                            order_text += "..."
                        st.code(order_text)
            else:
                st.success(f"✅ {total_zones} zones de texte détectées")

            # Stocker les résultats dans la session
            st.session_state.zone_results = zone_results

            # Affichage compact des zones détectées
            with st.expander(f"📊 Voir les {zone_results['total_zones']} zones détectées", expanded=False):
                col_img_annotated, col_zones_list = st.columns([1, 1])

                with col_img_annotated:
                    if zone_results["annotated_image"] and os.path.exists(zone_results["annotated_image"]):
                        st.image(zone_results["annotated_image"],
                                caption="Zones détectées",
                                use_column_width=True)

                with col_zones_list:
                    st.markdown("**Zones détectées:**")

                    # Affichage différencié selon le type de détection
                    if use_intelligent_detection:
                        # Affichage enrichi pour le système intelligent
                        for zone in zone_results["zones"][:5]:
                            zone_type = zone.get("type", "unknown")
                            type_emojis = {
                                "header": "🏷️", "price": "💰", "date": "📅",
                                "address": "🏠", "reference": "📄", "paragraph": "📝",
                                "signature": "✍️", "footer": "📋", "unknown": "❓"
                            }
                            emoji = type_emojis.get(zone_type, "📄")
                            confidence = zone.get("confidence", 0)
                            content = zone.get("content", "")
                            preview = content[:20] + "..." if len(content) > 20 else content

                            st.write(f"{emoji} **Zone {zone['zone_id']}** ({zone_type})")
                            st.caption(f"Confiance: {confidence:.1%} | '{preview}'")
                    else:
                        # Affichage classique
                        for zone in zone_results["zones"][:5]:
                            st.write(f"Zone {zone['zone_id']}: {zone['coordinates']['width']}×{zone['coordinates']['height']}px")

                    if len(zone_results["zones"]) > 5:
                        st.write(f"... et {len(zone_results['zones']) - 5} autres zones")

                    # Bouton de téléchargement ZIP compact
                    if zone_results["zones"]:
                        zip_path = os.path.join(zone_results["output_directory"], "zones_texte.zip")
                        with zipfile.ZipFile(zip_path, 'w') as zipf:
                            for zone in zone_results["zones"]:
                                if os.path.exists(zone["path"]):
                                    zipf.write(zone["path"], zone["filename"])

                            if zone_results["annotated_image"] and os.path.exists(zone_results["annotated_image"]):
                                zipf.write(zone_results["annotated_image"],
                                         os.path.basename(zone_results["annotated_image"]))

                        with open(zip_path, "rb") as f:
                            st.download_button(
                                "📦 Télécharger zones (ZIP)",
                                f.read(),
                                file_name="zones_texte.zip",
                                mime="application/zip"
                            )
        elif zone_results and not zone_results["success"]:
            st.warning(f"⚠️ Détection des zones échouée: {zone_results.get('error', 'Erreur inconnue')}")
            st.info("Le traitement OCR continuera sur l'image complète")

    # --- OCR principal selon la configuration
    st.markdown("### 📝 Reconnaissance de texte (OCR)")

    # Variables pour les résultats
    results = None
    best_method = None
    word_file = None
    zone_ocr_results = None

    # Déterminer le mode de traitement
    if selected_doc_type == "default":
        # Mode standard : OCR classique uniquement
        with st.spinner("🔄 Analyse OCR standard en cours..."):
            results, best_method, word_file = run_all_ocr_methods(image_paths)
        st.success("✅ Analyse OCR terminée !")

    else:
        # Mode avancé avec zones - OCR automatique sur les zones
        if zone_results and zone_results["success"]:
            zone_ocr_results = {}

            with st.spinner("🔄 Analyse OCR des zones en cours..."):
                for zone in zone_results["zones"]:
                    if os.path.exists(zone["path"]):
                        try:
                            zone_results_ocr, zone_best, _ = run_all_ocr_methods([zone["path"]])
                            zone_ocr_results[zone["zone_id"]] = {
                                "zone_info": zone,
                                "ocr_results": zone_results_ocr,
                                "best_method": zone_best,
                                "best_text": "\n".join(zone_results_ocr[zone_best]["lines"]),
                                "confidence": zone_results_ocr[zone_best]["avg_conf"]
                            }
                        except Exception as e:
                            zone_ocr_results[zone["zone_id"]] = {
                                "zone_info": zone,
                                "error": str(e)
                            }

            # Ajouter l'ordre de lecture intelligent aux résultats
            if "reading_order" in zone_results:
                zone_ocr_results["reading_order"] = zone_results["reading_order"]
            
            st.session_state.zone_ocr_results = zone_ocr_results

            # Compter les zones réussies (ignorer les clés spéciales)
            successful_zones = [z for z in zone_ocr_results.values() 
                               if isinstance(z, dict) and "error" not in z]
            if successful_zones:
                st.success(f"✅ OCR terminé sur {len(successful_zones)} zones")

                # Créer le document Word avec toutes les zones
                word_file = _create_zones_word_document(zone_ocr_results, image_paths[0])

                # Calculer les statistiques globales pour l'affichage
                total_confidence = sum(z["confidence"] for z in successful_zones)
                avg_confidence = total_confidence / len(successful_zones)
                total_lines = sum(len(z["ocr_results"][z["best_method"]]["lines"]) for z in successful_zones)

                # Créer un résultat synthétique pour l'affichage
                results = {
                    "zones_combined": {
                        "lines": [z["best_text"] for z in successful_zones if z["best_text"].strip()],
                        "avg_conf": avg_confidence,
                        "confs": [z["confidence"] for z in successful_zones]
                    }
                }
                best_method = "zones_combined"
                
                # Afficher les statistiques par moteur
                st.markdown("### 📊 Statistiques par moteur OCR")
                method_stats = {"tesseract": [], "easyocr": [], "doctr": []}
                
                for zone_data in successful_zones:
                    ocr_results = zone_data["ocr_results"]
                    for method, result in ocr_results.items():
                        if method in method_stats:
                            method_stats[method].append(result.get("avg_conf", 0))
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    if method_stats["tesseract"]:
                        avg_tesseract = sum(method_stats["tesseract"]) / len(method_stats["tesseract"])
                        st.metric("Tesseract", f"{avg_tesseract:.1f}%", f"{len(method_stats['tesseract'])} zones")
                    else:
                        st.metric("Tesseract", "0%", "0 zones")
                
                with col2:
                    if method_stats["easyocr"]:
                        avg_easyocr = sum(method_stats["easyocr"]) / len(method_stats["easyocr"])
                        st.metric("EasyOCR", f"{avg_easyocr:.1f}%", f"{len(method_stats['easyocr'])} zones")
                    else:
                        st.metric("EasyOCR", "0%", "0 zones")
                
                with col3:
                    if method_stats["doctr"]:
                        avg_doctr = sum(method_stats["doctr"]) / len(method_stats["doctr"])
                        st.metric("DocTR", f"{avg_doctr:.1f}%", f"{len(method_stats['doctr'])} zones")
                    else:
                        st.metric("DocTR", "0%", "0 zones")
            else:
                st.error("❌ Aucune zone n'a pu être traitée avec succès")
                
                # Fallback vers OCR standard
                st.info("🔄 Tentative de traitement OCR standard...")
                with st.spinner("🔄 Analyse OCR standard en cours..."):
                    results, best_method, word_file = run_all_ocr_methods(image_paths)
                st.success("✅ Analyse OCR standard terminée !")
        else:
            st.warning("⚠️ Aucune zone détectée - Passage en mode OCR standard")
            
            # Fallback vers OCR standard
            with st.spinner("🔄 Analyse OCR standard en cours..."):
                results, best_method, word_file = run_all_ocr_methods(image_paths)
            st.success("✅ Analyse OCR standard terminée !")

        # Affichage des résultats par zone
        if zone_ocr_results:
            with st.expander("🎯 Détails par zone", expanded=False):
                for zone_id, zone_data in zone_ocr_results.items():
                    # Ignorer les clés spéciales qui ne sont pas des zones
                    if not isinstance(zone_data, dict) or zone_id == "reading_order":
                        continue
                        
                    if "error" in zone_data:
                        st.error(f"Zone {zone_id}: {zone_data['error']}")
                        continue

                    best_method_zone = zone_data.get("best_method", "tesseract")
                    confidence = zone_data.get("confidence", 0)

                    with st.expander(f"📄 Zone {zone_id} - {best_method_zone.upper()} ({confidence:.1f}%)"):
                        col_preview, col_text = st.columns([1, 2])

                        with col_preview:
                            if os.path.exists(zone_data["zone_info"]["path"]):
                                st.image(zone_data["zone_info"]["path"],
                                       caption=f"Zone {zone_id}",
                                       use_column_width=True)

                        with col_text:
                            # Afficher tous les résultats OCR pour cette zone
                            st.markdown("**🔍 Résultats de tous les moteurs OCR :**")
                            
                            ocr_results = zone_data["ocr_results"]
                            for method, result in ocr_results.items():
                                method_name = method.upper()
                                avg_conf = result.get("avg_conf", 0)
                                lines = result.get("lines", [])
                                text = "\n".join(lines) if lines else "Aucun texte détecté"
                                
                                # Indiquer le meilleur résultat
                                is_best = method == best_method_zone
                                status = "🏆 MEILLEUR" if is_best else ""
                                
                                with st.expander(f"{method_name} ({avg_conf:.1f}%) {status}", expanded=is_best):
                                    st.text_area(
                                        f"Texte {method_name}",
                                        value=text,
                                        height=80,
                                        key=f"zone_{zone_id}_{method}_text",
                                        disabled=True
                                    )
                            
                            # Zone d'édition pour le meilleur résultat
                            st.markdown("**✏️ Édition du meilleur résultat :**")
                            edited_text = st.text_area(
                                f"Texte extrait (meilleur: {best_method_zone.upper()})",
                                value=zone_data["best_text"],
                                height=100,
                                key=f"zone_{zone_id}_text"
                            )

                            if st.button(f"💾 Sauvegarder", key=f"save_zone_{zone_id}"):
                                zone_text_dir = os.path.join("output", "zone_texts")
                                os.makedirs(zone_text_dir, exist_ok=True)

                                text_filename = f"zone_{zone_id}_{best_method_zone}.txt"
                                text_path = os.path.join(zone_text_dir, text_filename)

                                with open(text_path, "w", encoding="utf-8") as f:
                                    f.write(edited_text)

                                st.success(f"Sauvegardé: {text_filename}")

    # --- Affichage des résultats
    if results and best_method:
        st.markdown("---")
        st.markdown("### 📊 Résultats de l'analyse")

        col_img, col_text = st.columns([1, 2])
        with col_img:
            st.image(image_paths[0], caption="Image analysée")

        with col_text:
            # Affichage selon le mode
            if selected_doc_type == "default":
                # Mode standard : affichage classique
                st.markdown(f"**🏆 Meilleur résultat : {best_method.upper()}**")
                best_data = results[best_method]

                # Indicateur de confiance
                conf_value = best_data['avg_conf']
                if conf_value >= 85:
                    conf_emoji, conf_label = "🟢", "Excellente"
                elif conf_value >= 70:
                    conf_emoji, conf_label = "🟠", "Bonne"
                elif conf_value >= 50:
                    conf_emoji, conf_label = "🟠", "Moyenne"
                else:
                    conf_emoji, conf_label = "🔴", "Faible"

                st.markdown(f"{conf_emoji} **Confiance : {conf_value:.1f}% ({conf_label})**")

                # Texte principal
                best_text = "\n".join(best_data['lines'])
                st.text_area("Texte extrait", value=best_text, height=200, key="main_text")

            else:
                # Mode avancé : affichage des zones combinées
                if zone_ocr_results:
                    # Compter les zones réussies (ignorer les clés spéciales)
                    successful_zones = [z for z in zone_ocr_results.values() 
                                       if isinstance(z, dict) and "error" not in z]

                    st.markdown(f"**🎯 Résultat par zones : {len(successful_zones)} zones traitées**")

                    if successful_zones:
                        # Confiance moyenne
                        avg_conf = sum(z["confidence"] for z in successful_zones) / len(successful_zones)
                        if avg_conf >= 85:
                            conf_emoji, conf_label = "🟢", "Excellente"
                        elif avg_conf >= 70:
                            conf_emoji, conf_label = "🟠", "Bonne"
                        elif avg_conf >= 50:
                            conf_emoji, conf_label = "🟠", "Moyenne"
                        else:
                            conf_emoji, conf_label = "🔴", "Faible"

                        st.markdown(f"{conf_emoji} **Confiance moyenne : {avg_conf:.1f}% ({conf_label})**")

                        # Récupérer l'ordre de lecture intelligent depuis les résultats
                        reading_order = []
                        if hasattr(zone_ocr_results, 'get') and callable(zone_ocr_results.get):
                            # Si c'est un dict avec une clé reading_order
                            reading_order = zone_ocr_results.get("reading_order", [])
                        else:
                            # Essayer de récupérer depuis les métadonnées des zones
                            for zone_data in successful_zones:
                                if "reading_order" in zone_data:
                                    reading_order = zone_data["reading_order"]
                                    break
                        
                        # Si pas d'ordre de lecture, utiliser l'ordre des IDs
                        if not reading_order:
                            reading_order = sorted(zone_ocr_results.keys())
                        
                        # Réorganiser le texte selon l'ordre de lecture intelligent
                        reorganized_text = reorganize_text_by_reading_order(
                            {zone_id: zone_data for zone_id, zone_data in zone_ocr_results.items() 
                             if isinstance(zone_data, dict) and "error" not in zone_data},
                            reading_order
                        )

                        # Afficher le texte réorganisé
                        st.markdown("### 📄 Texte extrait (ordre de lecture intelligent)")
                        st.markdown("*Le texte est organisé selon l'ordre logique de lecture du document*")
                        
                        if reorganized_text and reorganized_text != "Aucun texte valide détecté dans les zones":
                            st.text_area(
                                "Texte réorganisé",
                                value=reorganized_text,
                                height=300,
                                key="reorganized_text"
                            )
                        else:
                            st.warning("⚠️ Aucun texte valide n'a pu être extrait des zones détectées")
                    else:
                        st.error("Aucune zone n'a pu être traitée avec succès")
                else:
                    st.warning("Aucun résultat OCR disponible")

        # Comparaison détaillée dans un expander (seulement pour le mode standard)
        if selected_doc_type == "default" and len(results) > 1:
            with st.expander("🔍 Comparaison détaillée des moteurs OCR", expanded=False):
                cols = st.columns(len(results))
                for i, (method, data) in enumerate(results.items()):
                    with cols[i]:
                        conf_value = data['avg_conf']

                        if conf_value >= 85:
                            conf_emoji = "🟢"
                        elif conf_value >= 70:
                            conf_emoji = "🟠"
                        elif conf_value >= 50:
                            conf_emoji = "🟠"
                        else:
                            conf_emoji = "🔴"

                        st.markdown(f"**{method.upper()}**")
                        st.markdown(f"{conf_emoji} {conf_value:.1f}%")

                        # Texte compact
                        method_text = "\n".join(data['lines'])
                        st.text_area(
                            f"Texte {method}",
                            value=method_text,
                            height=150,
                            key=f"text_{method}"
                        )

                        # Statistiques compactes
                        st.caption(f"📄 {len(data['lines'])} lignes • 📝 {sum(len(line) for line in data['lines'])} caractères")

        # --- Export et téléchargement
        if word_file and os.path.exists(word_file):
            st.markdown("---")
            st.markdown("### 📄 Export")

            # Export Word classique
            col_export, col_correction = st.columns(2)

            with col_export:
                if selected_doc_type == "default":
                    st.success(f"🏆 Meilleur résultat : {best_method.upper()}")
                    button_text = "📄 Télécharger le document Word"
                    
                    with open(word_file, "rb") as f:
                        st.download_button(
                            button_text,
                            f.read(),
                            file_name=os.path.basename(word_file),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    successful_zones = len([z for z in zone_ocr_results.values() if isinstance(z, dict) and "error" not in z]) if zone_ocr_results else 0
                    st.success(f"🎯 Document avec {successful_zones} zones")
                    
                    # Deux boutons d'export pour le mode intelligent
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        with open(word_file, "rb") as f:
                            st.download_button(
                                "📄 Document Word (zones détaillées)",
                                f.read(),
                                file_name=os.path.basename(word_file),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    
                    with col2:
                        # Créer un document Word simple avec juste le texte réorganisé
                        if st.button("📄 Créer document Word (texte réorganisé)", key="create_simple_word"):
                            simple_word_file = create_simple_word_document(zone_ocr_results, image_paths[0])
                            if simple_word_file and os.path.exists(simple_word_file):
                                with open(simple_word_file, "rb") as f:
                                    st.download_button(
                                        "📄 Télécharger Word (texte réorganisé)",
                                        f.read(),
                                        file_name=os.path.basename(simple_word_file),
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                            else:
                                st.error("❌ Erreur lors de la création du document simple")

            with col_correction:
                st.markdown("**📝 Correction manuelle**")
                corrected_file = st.file_uploader(
                    "Déposez votre Word corrigé",
                    type=["docx"],
                    key="corrected",
                    help="Téléversez un document Word corrigé pour améliorer l'apprentissage"
                )

                if corrected_file is not None:
                    os.makedirs("corrected", exist_ok=True)
                    doc = docx.Document(corrected_file)
                    corrected_text = "\n".join([p.text for p in doc.paragraphs])

                    base_name = uploaded_file.name if uploaded_file else "unknown"
                    txt_name = os.path.splitext(base_name)[0] + "_corrige.txt"
                    corrected_path = os.path.join("corrected", txt_name)

                    with open(corrected_path, "w", encoding="utf-8") as f:
                        f.write(corrected_text)
                    st.success(f"✅ Correction enregistrée")


